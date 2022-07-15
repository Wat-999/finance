#先通过一段演示代码来感受下python-docx库的使用效果。
import docx
file = docx.Document()  #在后台创建一个word文档对象

file.add_paragraph('螃蟹在剥我的壳，笔记本在写我')  #通过file.add_paragraph添加段落
file.add_paragraph('漫天的我落在枫叶上雪花上')
file.add_paragraph('而你在想我')

file.save('/Users/macbookair/Desktop/简历/三行情书.docx')  #file.save()函数命名并保存word文档，可绝对位置可相对位置
print('wrod文档生成完毕')

#python-docx库的基本操作

#创建、打开、及保存word文档
#通过docx.Document()函数可以创建一个word文档
import docx
file = docx.Document()

#如果想打开已经存在的word文档，只要在Document后的括号里填写word文档的路径即可
file= docx.Document('/Users/macbookair/Desktop/简历/三行情书.docx')

#通过save()函数可以命名并保存word文档
file.save('/Users/macbookair/Desktop/简历/三行情书.docx')
print('三行情书2生成完毕')

#2添加标题：通过add_heading()函数可以添加标题
file.add_heading('三行情书2', level=0)  #level=0表示标题级别为0，字号较大。这里不推荐采用该方式添加标题，因为它会默认添加下划线
#可以用下面将的添加段落的方式来创建标题

#添加段落： add_paragraph()可以添加段落
file.add_paragraph('我喜欢你')
file.add_paragraph('上一句话是假的')
file.add_paragraph('上一句话也是假的')

#添加图片  #add_picture()函数可以添加图片
file.add_picture('/Users/macbookair/Desktop/简历/水墨.png')

#添加分页符: add_page_break()函数可以添加分页符，实现分页
file.add_page_break()

#添加表格：add_table()函数可以添加表格；默认的表格样式没有边框
table = file.add_table(rows=1, cols=3)  #添加一个1行3列的表格 rows参数指定表格的行数，cols参数指定表格的列数
table.cell(0, 0).text = '克制'  #设置第1行1列单元格的内容 cell(a, b)中，a表示第几行，b表示第几列
table.cell(0, 1).text = '再克制'  #设置第1行2列单元格的内容
table.cell(0, 2).text = '在吗'  #设置第1行3列单元格的内容

#读取word文档的内容
file = docx.Document('/Users/macbookair/Desktop/简历/三行情书.docx') #打开word文档
for paragraph in file.paragraph:
    print(paragraph.text)     #打印输出各段落文本，注意不要漏写.text
    #代码中的file.paragraph即该word文档中所有的段落，通过for循环即可打印输出每个段落

#所有代码汇总如下：
# 1.创建Word对象
import docx

file = docx.Document()
# 2.添加标题
file.add_heading('三行情书2', level=0)
# 3.添加段落文字
file.add_paragraph('我喜欢你')
file.add_paragraph('上一句话是假的')
file.add_paragraph('上一句话也是假的')
# 4.添加图片
file.add_picture('水墨.png')  # 需要你自己设置一个图片地址
# 5.添加分页符
file.add_page_break()
# 6.添加表格
table = file.add_table(rows=1, cols=3)
table.cell(0, 0).text = '克制'
table.cell(0, 1).text = '再克制'
table.cell(0, 2).text = '在吗'
# 7.文档保存，存储文件夹需提前创建
file.save('三行情书2.docx')  # 这是个相对路径写法（代码所在文件夹），也可以写成绝对路径：E:\\三行情书2.docx
print('三行情书2生成完毕')



