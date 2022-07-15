
import docx  # 下面的import的代码其实都可以写到这个import下面
file = docx.Document()

# # 1.设置中文字体
from docx.oxml.ns import qn
file.styles['Normal'].font.name = u'微软雅黑'  # 可换成word里面任意字体
file.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')  # 这边记得也得填一下字体名称

# 写入若干段落
# # 2.设置字体颜色和大小
p = file.add_paragraph()   #添加段落
run = p.add_run('螃蟹在剥我的壳，笔记本在写我')  #通过add_run()函数可以在一个段落的末尾新增内容
font = run.font  #利用run.font便可以调用它的字体属性
#这里的run可以裂解为段落里的文字，只有段落里的文字可以设置字体属性，通过file.add_paragraph()生成的段落是无法直接设置字体的，
#所以不能直接写font = p.font

from docx.shared import Pt
font.size = Pt(26)  #利用font.size设置字体大小，其中pt()里的数字即要设置的字体大小，单位为"磅"
from docx.shared import RGBColor
font.color.rgb = RGBColor(54, 95, 145)  #利用font属性可以修改字体颜色
#上述代码在设置颜色时使用了RGB颜色模式，它通过混合R(红)G(绿)B(蓝)三种原色来得到各种颜色，这里的(54, 95, 145)是一种蓝色系的颜色。
#如果想查看更多颜色的RGB色值，可以利用搜索引擎搜索"RGB颜色对照表"

# # 3.设置字体粗体、斜体和下划线：font属性还可以设置字体的粗体、下体和下画线效果
p = file.add_paragraph()
run = p.add_run('漫天的我落在枫叶上雪花上')
font = run.font
font.bold = True  # 粗体
font.italic = True  # 斜体
font.underline = True  # 下划线

# # 4.设置居中对齐
from docx.enum.text import WD_ALIGN_PARAGRAPH
p = file.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  #两端对齐 居中对齐：WD_ALIGN_PARAGRAPH.CENTER
p.add_run('而你在想我')

# #  5.设置首行缩进
from docx.shared import Inches
p = file.add_paragraph()
p.paragraph_format.first_line_indent = Inches(0.32)  #Inches(0.32)括号里的数字为缩进量，单位为尺寸
#假设这个段落文本的字体大小是小四号，从之前的字体大小单位对照表可查到，小四号对应0.16寸，而缩进一般是2个字符长度，所以设置0。32英寸的首行缩进量
#还可以对其进行适当微调。对于其他大小的字体，可以参照此方法计算首行缩进量
p.add_run('设置首行缩进示例文字')

# # 6.设置行距
from docx.shared import Pt
p = file.add_paragraph()
p.paragraph_format.line_spacing = Pt(16)  # 行距，16磅对应三号字体大小
p.add_run('设置行距示例文字')

# #7.设置段前距和段后距
from docx.shared import Pt
p = file.add_paragraph()
p.paragraph_format.space_before = Pt(14)  # 段前距,14磅对应4号字体大小
p.paragraph_format.space_after = Pt(14)  # 段后距
p.add_run('设置段前段后距示例文字')

# # 8.设置段落序号
file.add_paragraph('项目符号示例文字', style='List Bullet')  #点序号
file.add_paragraph('数字编号示例文字', style='List Number')  #数字序号

# # 9.设置表格
table = file.add_table(rows=2, cols=3, style='Light Shading Accent 1')  #style参数设置表格样式
table.cell(0, 0).text = '第一句'  # 第一行第一列
table.cell(0, 1).text = '第二句'  # 第一行第二列
table.cell(0, 2).text = '第三句'  # 第一行第三列
table.cell(1, 0).text = '克制'  # 第二行第一列
table.cell(1, 1).text = '再克制'  # 第二行第二列
table.cell(1, 2).text = '"在吗"'  # 第三行第三列

# # 10.设置图表
file.add_picture('/Users/macbookair/Desktop/简历/水墨.png', width=Inches(3), height=Inches(3))
#width为指定宽，height为指定高， Inches即英寸，1英寸等于2.54.cm
last_paragraph = file.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 保存，得首先创建出保存文件夹，且Word文件不要打开
file.save("三行情书.docx")  # 这是个相对路径写法（代码所在文件夹），也可以写成绝对路径：E:\\三行情书2.docx
print('Word生成完毕！')